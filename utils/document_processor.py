"""
Document Processor for AURA
Processes various document types (PDF, DOCX, TXT, images)
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import io
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image

# Make pytesseract optional
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    
from utils.logger import logger
from utils.supabase_client import supabase_client
from utils.pinecone_store import pinecone_store
from utils.embedding_generator import embedding_generator

class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    def __init__(self):
        """Initialize document processor"""
        self.supported_extensions = ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg']
        logger.info(f"Document processor initialized. Supported formats: {self.supported_extensions}")
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document file and extract text
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return {"success": False, "error": "File not found"}
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            logger.error(f"Unsupported file format: {extension}")
            return {"success": False, "error": f"Unsupported format: {extension}"}
        
        logger.info(f"Processing file: {file_path.name} ({extension})")
        
        try:
            if extension == '.pdf':
                return self._process_pdf(file_path)
            elif extension == '.docx':
                return self._process_docx(file_path)
            elif extension == '.txt':
                return self._process_txt(file_path)
            elif extension in ['.png', '.jpg', '.jpeg']:
                return self._process_image(file_path)
            else:
                return {"success": False, "error": "Format not implemented"}
                
        except Exception as e:
            logger.error(f"Error processing file {file_path.name}: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(str(file_path))
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            metadata = {
                "filename": file_path.name,
                "format": "pdf",
                "pages": len(reader.pages),
                "text_length": len(text)
            }
            
            logger.info(f"Extracted {len(text)} characters from {len(reader.pages)} pages")
            
            return {
                "success": True,
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = Document(str(file_path))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            metadata = {
                "filename": file_path.name,
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "text_length": len(text)
            }
            
            logger.info(f"Extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs")
            
            return {
                "success": True,
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                "filename": file_path.name,
                "format": "txt",
                "text_length": len(text)
            }
            
            logger.info(f"Extracted {len(text)} characters from text file")
            
            return {
                "success": True,
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing TXT: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            
            # Try OCR if available
            if TESSERACT_AVAILABLE:
                try:
                    text = pytesseract.image_to_string(image)
                    logger.info("OCR extraction successful")
                except Exception as ocr_error:
                    logger.warning(f"OCR failed: {ocr_error}")
                    text = f"[Image: {file_path.name} - OCR extraction failed]"
            else:
                logger.warning("Tesseract OCR not available. Install pytesseract for text extraction.")
                text = f"[Image: {file_path.name} - OCR not available]"
            
            metadata = {
                "filename": file_path.name,
                "format": file_path.suffix[1:],
                "size": image.size,
                "mode": image.mode,
                "text_length": len(text),
                "ocr_available": TESSERACT_AVAILABLE
            }
            
            return {
                "success": True,
                "text": text.strip() if text else "",
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into chunks for better embedding and retrieval
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
    
    def process_and_store(self, file_path: str, document_id: Optional[str] = None) -> Dict[str, Any]:
        """
        Process document, upload to Supabase, and store embeddings in Pinecone
        
        Args:
            file_path: Path to the document
            document_id: Optional custom document ID
            
        Returns:
            Processing results
        """
        file_path = Path(file_path)
        
        # Process document
        result = self.process_file(file_path)
        
        if not result.get("success"):
            return result
        
        text = result["text"]
        metadata = result["metadata"]
        
        # Generate document ID
        if not document_id:
            document_id = f"{file_path.stem}_{hash(text) % 10000}"
        
        try:
            # Upload file to Supabase
            logger.info("Uploading file to Supabase...")
            supabase_path = f"documents/{file_path.name}"
            public_url = supabase_client.upload_file(
                str(file_path),
                supabase_path,
                content_type=self._get_content_type(file_path.suffix)
            )
            
            if public_url:
                metadata['supabase_url'] = public_url
                logger.info(f"File uploaded to Supabase: {public_url}")
            
            # Chunk text for better retrieval
            chunks = self.chunk_text(text)
            
            # Store each chunk in Pinecone
            logger.info(f"Storing {len(chunks)} chunks in Pinecone...")
            stored_count = 0
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"{document_id}_chunk_{i}"
                chunk_metadata = {
                    **metadata,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
                
                success = pinecone_store.upsert_document(chunk_id, chunk, chunk_metadata)
                if success:
                    stored_count += 1
            
            logger.info(f"Successfully stored {stored_count}/{len(chunks)} chunks")
            
            return {
                "success": True,
                "document_id": document_id,
                "text_length": len(text),
                "chunks_stored": stored_count,
                "supabase_url": public_url,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error in process_and_store: {str(e)}")
            return {"success": False, "error": str(e)}
    
    def _get_content_type(self, extension: str) -> str:
        """Get MIME type for file extension"""
        content_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg'
        }
        return content_types.get(extension.lower(), 'application/octet-stream')

# Global instance
document_processor = DocumentProcessor()
